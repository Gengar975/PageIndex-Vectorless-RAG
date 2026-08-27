from pathlib import Path
from typing import Any
import re

import fitz
from docx import Document
from openpyxl import load_workbook


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx"}


def discover_documents(path: str | Path) -> list[Path]:
    path = Path(path)

    if path.is_file():
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {path.suffix}"
            )
        return [path]

    if not path.exists():
        raise FileNotFoundError(path)

    return sorted(
        p for p in path.rglob("*")
        if p.is_file()
        and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def parse_document(path: str | Path) -> dict[str, Any]:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf(path)

    if suffix == ".docx":
        return extract_docx(path)

    if suffix == ".xlsx":
        return extract_xlsx(path)

    raise ValueError(f"Unsupported document type: {suffix}")


def extract_pdf(path: Path) -> dict[str, Any]:
    pages = []

    with fitz.open(path) as pdf:
        for number, page in enumerate(pdf, start=1):
            text = page.get_text("text")

            pages.append({
                "page_id": f"page_{number:03d}",
                "page_number": number,
                "type": "page",
                "title": f"Page {number}",
                "text": text.strip(),
            })

    return {
        "document_name": path.name,
        "source": str(path),
        "file_type": "pdf",
        "pages": pages,
    }


def extract_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)

    pages = []
    current = []
    page_number = 1

    def flush():
        nonlocal current, page_number

        if not current:
            return

        pages.append({
            "page_id": f"page_{page_number:03d}",
            "page_number": page_number,
            "type": "page",
            "title": f"Page {page_number}",
            "text": "\n".join(
                item["text"]
                for item in current
            ).strip(),
            "blocks": current,
        })

        page_number += 1
        current = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:
            style_name = ""

            try:
                style_name = paragraph.style.name or ""
            except Exception:
                style_name = ""

            is_heading = (
                style_name.lower().startswith("heading")
            )

            current.append({
                "text": text,
                "style": style_name,
                "is_heading": is_heading,
            })

        # DOCX does not expose rendered page numbers reliably.
        # Explicit page breaks are treated as logical page boundaries.
        xml = paragraph._p.xml

        if (
            'w:type="page"' in xml
            or "w:type='page'" in xml
        ):
            flush()

    # Preserve tables.
    for table in doc.tables:
        rows = []

        for row in table.rows:
            rows.append(
                " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                )
            )

        if rows:
            table_text = "\n".join(rows)

            current.append({
                "text": table_text,
                "style": "Table",
                "is_heading": False,
            })

    if current or not pages:
        flush()

    return {
        "document_name": path.name,
        "source": str(path),
        "file_type": "docx",
        "pages": pages,
    }


def extract_xlsx(path: Path) -> dict[str, Any]:
    workbook = load_workbook(
        filename=path,
        data_only=False,
        read_only=True,
    )

    pages = []

    for index, worksheet in enumerate(
        workbook.worksheets,
        start=1,
    ):
        rows = []

        for row in worksheet.iter_rows(values_only=True):
            values = [
                "" if value is None else str(value)
                for value in row
            ]

            while values and values[-1] == "":
                values.pop()

            if values:
                rows.append(" | ".join(values))

        text = "\n".join(rows).strip()

        pages.append({
            "page_id": f"worksheet_{index:03d}",
            "page_number": index,
            "type": "worksheet",
            "title": worksheet.title,
            "worksheet_name": worksheet.title,
            "text": text,
        })

    workbook.close()

    return {
        "document_name": path.name,
        "source": str(path),
        "file_type": "xlsx",
        "pages": pages,
    }
